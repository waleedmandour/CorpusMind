"""
File-format ingestion (§8.1).

Each parser returns the cleaned plain text. Supported formats:
  - .txt   (auto-detected encoding via charset-normalizer)
  - .docx  (python-docx)
  - .pdf   (pypdf)
  - .html  (BeautifulSoup)
  - .xml   (lxml — text content extracted)
  - .csv   (one column or a 'text' column)
  - .md    (treated as text after stripping simple markdown)
"""
from __future__ import annotations

import csv
import io
from pathlib import Path

from bs4 import BeautifulSoup
from charset_normalizer import from_bytes

from app.logging import get_logger

log = get_logger(__name__)


SUPPORTED_FORMATS = {"txt", "docx", "pdf", "html", "htm", "xml", "csv", "md"}


def detect_format(filename: str) -> str:
    ext = Path(filename).suffix.lower().lstrip(".")
    if ext == "htm":
        ext = "html"
    if ext not in SUPPORTED_FORMATS:
        raise ValueError(f"Unsupported file format: .{ext} (supported: {sorted(SUPPORTED_FORMATS)})")
    return ext


def decode_bytes(raw: bytes) -> tuple[str, str]:
    """Detect encoding and return (text, encoding_name)."""
    result = from_bytes(raw).best()
    if result is None:
        # Last-resort UTF-8 with replacement chars
        return raw.decode("utf-8", errors="replace"), "utf-8-replacement"
    return str(result), result.encoding or "utf-8"


def _clean_text(text: str) -> str:
    """Minimal cleaning: normalize whitespace, strip control chars.
    Boilerplate stripping is intentionally conservative — we don't want to
    silently drop legitimate content (§8.1's "visible, inspectable" principle)."""
    # Normalize Unicode line endings + collapse runs of whitespace
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Strip zero-width chars that can silently disagree between tools (§8.1)
    text = text.replace("\u200b", "").replace("\u200c", "").replace("\u200d", "")
    text = text.replace("\ufeff", "")  # BOM
    # Collapse 3+ newlines to 2
    while "\n\n\n" in text:
        text = text.replace("\n\n\n", "\n\n")
    return text.strip()


def parse_txt(raw: bytes) -> str:
    text, _ = decode_bytes(raw)
    return _clean_text(text)


def parse_md(raw: bytes) -> str:
    text, _ = decode_bytes(raw)
    # Conservative markdown strip: headers, emphasis, code, links
    out_lines = []
    in_code = False
    for line in text.splitlines():
        if line.strip().startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            out_lines.append(line)
            continue
        # Strip leading # for headers
        if line.lstrip().startswith("#"):
            line = line.lstrip("#").lstrip()
        # Strip bold/italic markers (very basic)
        for marker in ("**", "__", "*", "_", "`"):
            line = line.replace(marker, "")
        # Strip link syntax: [text](url) -> text
        if "](" in line and line.find("]") < line.find("]("):
            import re
            line = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", line)
        out_lines.append(line)
    return _clean_text("\n".join(out_lines))


def parse_docx(raw: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(raw))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also pull table cell text
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return _clean_text("\n\n".join(parts))


def parse_pdf(raw: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(raw))
    parts = []
    for page in reader.pages:
        t = page.extract_text() or ""
        if t.strip():
            parts.append(t)
    return _clean_text("\n".join(parts))


def parse_html(raw: bytes) -> str:
    text, _enc = decode_bytes(raw)
    soup = BeautifulSoup(text, "lxml")
    # Strip script/style
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    # Prefer <article>, <main>, or fall back to whole body
    for candidate in ("article", "main", "body"):
        node = soup.find(candidate)
        if node and node.get_text(strip=True):
            return _clean_text(node.get_text(separator="\n"))
    return _clean_text(soup.get_text(separator="\n"))


def parse_xml(raw: bytes) -> str:
    text, _ = decode_bytes(raw)
    soup = BeautifulSoup(text, "lxml-xml")
    # Take all leaf text nodes
    return _clean_text(soup.get_text(separator="\n"))


def parse_csv(raw: bytes) -> str:
    text, _enc = decode_bytes(raw)
    reader = csv.DictReader(io.StringIO(text))
    # Prefer a column literally named 'text', otherwise the first column
    field = "text" if "text" in (reader.fieldnames or []) else (reader.fieldnames or ["text"])[0]
    parts = [row[field] for row in reader if row.get(field, "").strip()]
    return _clean_text("\n".join(parts))


_PARSERS = {
    "txt": parse_txt,
    "md": parse_md,
    "docx": parse_docx,
    "pdf": parse_pdf,
    "html": parse_html,
    "xml": parse_xml,
    "csv": parse_csv,
}


def parse_file(filename: str, raw: bytes) -> tuple[str, str, str]:
    """Parse a file's raw bytes → (cleaned_text, format, detected_encoding)."""
    fmt = detect_format(filename)
    if fmt not in _PARSERS:
        raise ValueError(f"No parser for format: {fmt}")
    # For binary formats (docx, pdf) the parser handles its own decoding.
    if fmt in {"docx", "pdf"}:
        text = _PARSERS[fmt](raw)
        encoding = "binary-decoded"
    else:
        text, encoding = decode_bytes(raw)
        text = _clean_text(text) if fmt in {"txt"} else _PARSERS[fmt](raw)
    log.info("ingest_parsed", filename=filename, format=fmt, encoding=encoding, chars=len(text))
    return text, fmt, encoding
